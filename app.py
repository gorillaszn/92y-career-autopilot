import streamlit as st
import google.generativeai as genai
import pypdf
import docx
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import json
import time
import zipfile
from datetime import datetime


# ============================================================
# 1. PAGE CONFIG & SESSION STATE
# ============================================================

st.set_page_config(page_title="92Y Career Auto-Pilot", page_icon="*", layout="wide")

STATE_DEFAULTS = {
    "model": None,
    "keywords": None,
    "keywords_confirmed": False,
    "match_score": None,
    "resume_md": None,
    "cover_letter_md": None,
    "interview_md": None,
    "generation_complete": False,
    "company_name": None,
    "generated_at": None,
    "ats_analysis": None,
}
for k, v in STATE_DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ============================================================
# 2. MODEL INITIALIZATION
# ============================================================

def init_model(api_key: str):
    genai.configure(api_key=api_key)
    try:
        for m in genai.list_models():
            if "flash" in m.name:
                return genai.GenerativeModel(m.name)
        for m in genai.list_models():
            if "pro" in m.name:
                return genai.GenerativeModel(m.name)
    except Exception:
        pass
    return genai.GenerativeModel("gemini-1.5-flash")


# ============================================================
# 3. FILE READER
# ============================================================

def read_file(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        reader = pypdf.PdfReader(uploaded_file)
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("PDF appears to be scanned/image-only.")
        return text
    elif name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        text = "\n".join([p.text for p in doc.paragraphs]).strip()
        if not text:
            raise ValueError("DOCX file appears empty.")
        return text
    raise ValueError(f"Unsupported file type: {uploaded_file.name}")


# ============================================================
# 4. LLM CALL WRAPPER
# ============================================================

def call_model(model, prompt: str, retries: int = 2) -> str:
    last_err = None
    for attempt in range(retries + 1):
        try:
            response = model.generate_content(prompt)
            text = response.text.strip()
            if text.startswith("```"):
                text = re.sub(r"^```(?:json)?\s*", "", text)
                text = re.sub(r"\s*```$", "", text)
            return text
        except Exception as e:
            last_err = e
            if attempt < retries:
                time.sleep(1.5 * (attempt + 1))
    raise last_err


# ============================================================
# 5. KNOWLEDGE BASE
# ============================================================

TRANSLATION_MAP = {
    "Property Book": "Capital Asset Portfolio ($15M+)",
    "Hand Receipt": "Custodial Asset Transfer Protocol",
    "GCSS-Army": "SAP ERP Systems",
    "FLIPL": "Forensic Financial Audit",
    "PLL": "Preventive Maintenance Logistics Program",
    "CIF": "Central Inventory & Distribution Facility",
    "CSDP": "Command Supply Discipline Program / Internal Compliance Audit",
    "PBUSE": "Automated Asset Tracking Systems",
    "NCOER": "Performance Evaluation / Annual Review",
    "TA-50": "Individual Equipment Accountability Program",
    "SSA": "Supply Support Activity / Regional Distribution Hub",
    "UBL": "Unit Basic Load / Critical Stock Reserve",
    "Class I": "Subsistence & Perishable Inventory",
    "Class II": "Administrative & General Supplies",
    "Class IV": "Construction & Barrier Materials",
    "Class IX": "Repair Parts & Supply Chain Maintenance",
    "S4 Shop": "Logistics Operations Center",
    "Motor Pool": "Fleet Maintenance Facility",
    "Battalion": "Regional Business Unit",
    "Brigade": "Divisional Headquarters",
    "Company Commander": "Operations Director",
}

GHOSTWRITER = {
    "E-4": [
        "Data Accuracy & Record Integrity",
        "Technical Execution of Standard Operating Procedures",
        "Inventory Control & Cycle Count Operations",
        "ERP Data Entry & Transaction Processing",
    ],
    "E-5": [
        "Team Leadership (10-20 personnel)",
        "Training Program Development",
        "Risk Assessment & Mitigation",
        "Budget Oversight ($2M-$5M)",
        "Customer/Stakeholder Liaison",
    ],
    "E-6": [
        "Team Leadership (20+ personnel)",
        "Training Program Development & Execution",
        "Operational Risk Management",
        "Budget Administration ($5M-$10M)",
        "Cross-Functional Stakeholder Management",
        "Process Improvement Initiatives",
        "Customer Service & Internal Liaison",
    ],
    "E-7": [
        "Strategic Planning & Organizational Oversight",
        "Audit & Compliance Program Management",
        "Policy Development & Implementation",
        "Budget Authority ($10M-$15M+)",
        "Senior Stakeholder Advisory",
        "Workforce Development Strategy",
    ],
    "E-8": [
        "Enterprise-Level Strategic Operations",
        "Inspector General-Level Audit Oversight",
        "Organizational Policy Architecture",
        "Executive Budget Authority ($15M+)",
        "C-Suite Advisory & Cross-Org Coordination",
    ],
}

INDUSTRY_TONE = {
    "Corporate (General)": (
        "Use business-neutral corporate language. Emphasize ROI, cost savings, efficiency, "
        "and operational excellence. Avoid military jargon entirely. "
        "If the role is Entry/Mid-level (Buyer, Specialist, Coordinator), use OPERATIONAL verbs: "
        "'Executed,' 'Processed,' 'Maintained,' 'Resolved.' "
        "If the role is Senior (Director, VP), use STRATEGIC verbs: 'Directed,' 'Spearheaded,' 'Optimized.'"
    ),
    "Defense Contractor": (
        "Use defense/aerospace industry language. Reference security clearances, ITAR, DFARS, "
        "government contracts, and controlled inventory. Emphasize Warfighter Readiness, "
        "Production Speed, Mission Assurance, and Accelerated Procurement per 2026 Executive Order priorities. "
        "Military familiarity is expected but translate MOS-specific jargon."
    ),
    "Federal (USAJOBS)": (
        "Use federal resume conventions. Be VERBOSE and THOROUGH. Include '40 Hours/Week' for each position. "
        "Use KSA-style detail. Match OPM qualification standards language. "
        "Include supervisor name/phone placeholders. This resume should be 3-5 pages, not 1-2."
    ),
    "Tech / SaaS": (
        "Use tech industry language. Emphasize data-driven decisions, automation, scalability, "
        "agile methodology, cross-functional collaboration, and customer success. "
        "Frame supply chain as 'Operations Management.' Keep it modern and concise."
    ),
}


# ============================================================
# 6. PROMPT BUILDERS
# ============================================================

def _project_header(industry):
    headers = {
        "Defense Contractor": "KEY MILITARY PROJECTS",
        "Corporate (General)": "KEY STRATEGIC INITIATIVES",
        "Federal (USAJOBS)": "RELEVANT PROJECT EXPERIENCE",
        "Tech / SaaS": "MAJOR OPERATIONS PROJECTS",
    }
    return headers.get(industry, "KEY PROJECTS")


def _contact_block(contact_info):
    name = contact_info.get("name", "[Full Name]")
    city = contact_info.get("city") or "[City, State]"
    phone = contact_info.get("phone") or "[Phone]"
    email = contact_info.get("email") or "[Email]"
    linkedin = contact_info.get("linkedin", "")
    parts = [f"**{city}**", f"**{phone}**", f"**{email}**"]
    if linkedin:
        parts.append(f"**{linkedin}**")
    return name, " | ".join(parts)


def _gap_statement(gap_info):
    if not gap_info or not gap_info.get("has_gap"):
        return ""
    gap_start = gap_info.get("start", "")
    gap_end = gap_info.get("end", "Present")
    activities = gap_info.get("activities", [])
    if not activities:
        activities = ["professional development", "skills training", "certification pursuit"]
    if len(activities) > 1:
        act_str = ", ".join(activities[:-1]) + f", and {activities[-1]}"
    else:
        act_str = activities[0]
    return (
        f"\n\nEMPLOYMENT GAP NOTE: The candidate has a gap from {gap_start} to {gap_end}. "
        f"Frame this positively as a 'Professional Development Period' focused on {act_str}. "
        f"Integrate this naturally into the Professional Summary or Experience section. "
        f"Do NOT draw attention to the gap. Present it as intentional career investment."
    )


def _context_block(rank, years, industry, target_title, keywords, user_data, contact_info=None, gap_info=None):
    rank_code = rank.split(" ")[0]
    ghost = GHOSTWRITER.get(rank_code, GHOSTWRITER["E-5"])
    trans = "\n".join(f"  - {k} -> {v}" for k, v in TRANSLATION_MAP.items())
    gh = "\n".join(f"  - {s}" for s in ghost)
    kw = "\n".join(f"  {i+1}. {k}" for i, k in enumerate(keywords))
    contact_str = ""
    if contact_info:
        name, contact_line = _contact_block(contact_info)
        contact_str = f"\n\nCANDIDATE CONTACT INFO (use exactly as provided):\n  Name: {name}\n  Contact Line: {contact_line}"
    gap_str = _gap_statement(gap_info)
    return f"""
CANDIDATE PROFILE:
  Rank: {rank}
  Years of Service: {years}
  Raw Experience Data:
  ---
  {user_data[:4000]}
  ---
{contact_str}
{gap_str}

TARGET POSITION:
  Title: {target_title}
  Industry: {industry}
  Industry Tone: {INDUSTRY_TONE.get(industry, INDUSTRY_TONE["Corporate (General)"])}

CONFIRMED JD KEYWORDS TO MIRROR (address ALL of these):
{kw}

MILITARY-TO-CIVILIAN TRANSLATIONS:
{trans}

GHOSTWRITER INFERRED SKILLS FOR {rank_code} (use if candidate data is thin or missing a JD requirement):
{gh}
"""


def prompt_keywords(job_desc):
    return f"""You are an expert ATS analyst.
Extract the 10-15 most important requirements, skills, and keywords from this job description.
Include BOTH hard skills AND soft skills (like Customer Service, Communication, Analytical Skills).
Prioritize skills that appear multiple times or are listed under "Required" / "Must Have."
Return ONLY a JSON array of strings. No preamble, no markdown fences, no explanation.
Order from most critical to least.
Example: ["Supply Chain Management", "SAP ERP", "Customer Service", "Vendor Negotiation"]
JOB DESCRIPTION:
{job_desc}"""


def prompt_company_extract(job_desc):
    return f"""Extract the company or organization name from this job description.
Return ONLY the company name as a plain string. No quotes, no explanation, no preamble.
If you cannot determine the company name, return exactly: Unknown Company
JOB DESCRIPTION:
{job_desc[:2000]}"""


def prompt_match_score(rank, years, keywords, user_data, target_title):
    rank_code = rank.split(" ")[0]
    ghost = GHOSTWRITER.get(rank_code, GHOSTWRITER["E-5"])
    gh = "\n".join(f"  - {s}" for s in ghost)
    kw = "\n".join(f"  - {k}" for k in keywords)
    return f"""You are a career match analyst for military veterans transitioning to civilian roles.
CANDIDATE:
  Rank: {rank}
  Years of Service: {years}
  Experience Data: {user_data[:3000] if user_data else "[No resume provided. Using rank-based inference only.]"}
  Standard skills for {rank_code} 92Y (Unit Supply Specialist):
{gh}
TARGET ROLE: {target_title}
JD KEYWORDS (the requirements this role demands):
{kw}
TASK: Evaluate how well this candidate matches the target role. For EACH keyword, determine if the candidate has it (from their data or from standard {rank_code} duties).
Return ONLY valid JSON with this exact structure. No preamble, no markdown fences:
{{
  "score": <integer 0-100>,
  "matched": ["keyword1", "keyword2"],
  "gaps": ["keyword3", "keyword4"],
  "summary": "<2 sentence assessment. Be direct about strengths and gaps.>"
}}
SCORING GUIDE:
- 80-100: Strong match. Most keywords covered by experience or rank duties.
- 60-79: Solid match. Some gaps but transferable skills fill them.
- 40-59: Stretch role. Multiple hard-skill gaps. Upskilling needed.
- Below 40: Significant mismatch. Major requirements missing."""


def prompt_resume(rank, years, industry, target_title, keywords, user_data, contact_info=None, gap_info=None, max_pages="2 pages (recommended)"):
    ctx = _context_block(rank, years, industry, target_title, keywords, user_data, contact_info, gap_info)
    if "1 page" in max_pages:
        page_rule = """PAGE LENGTH: STRICT 1 PAGE. This is non-negotiable.
- Professional Summary: 2 sentences max.
- Core Competencies: 4 items max.
- Projects section: OMIT entirely.
- Professional Experience: 3-4 bullets max. Pick only the highest-impact ones that match the JD.
- Education: 1 line.
Keep every bullet under 2 lines. Cut anything that doesn't directly match a JD keyword."""
    elif "3+" in max_pages:
        page_rule = """PAGE LENGTH: 3-5 pages. Use verbose federal resume style.
- Include full KSA detail, 40 Hours/Week notations, supervisor placeholders.
- Expand each bullet into 2-3 sentences with full STAR context.
- Add additional duty positions and collateral assignments."""
    else:
        page_rule = """PAGE LENGTH: STRICT 2 PAGES MAX. This is non-negotiable.
- If content exceeds 2 pages, cut the least relevant bullets first.
- Prioritize bullets that directly match JD keywords.
- Keep Professional Summary to 3 sentences.
- Core Competencies: 5-6 items max.
- Professional Experience: 5 bullets max. Each bullet should be 1-2 lines.
- Compact is better than comprehensive. Recruiters spend 7 seconds on a resume."""
    if contact_info and contact_info.get("name"):
        name = contact_info["name"]
        city = contact_info.get("city") or "[City, State]"
        phone = contact_info.get("phone") or "[Phone]"
        email = contact_info.get("email") or "[Email]"
        linkedin = contact_info.get("linkedin", "")
        contact_parts = [f"**{city}**", f"**{phone}**", f"**{email}**"]
        if linkedin:
            contact_parts.append(f"**{linkedin}**")
        header_block = f"# {name}\n### **{target_title}**\n{' | '.join(contact_parts)}"
    else:
        header_block = f'# [Candidate Name]\n### **{target_title}**\n**[City, State]** | **[Phone]** | **[Email]** | **[LinkedIn URL]**'
    proj_header = _project_header(industry)
    return f"""You are a Career Architect for U.S. Army 92Y veterans.
{ctx}

{page_rule}

RESUME RULES:
1. THE NO-REPEAT PROTOCOL:
   - CORE COMPETENCIES: Short keyword phrases with brief context.
     Example: "Vendor Negotiation: Managed $5M contracts across 15 suppliers."
   - PROFESSIONAL EXPERIENCE: Detailed STAR-method bullets with different wording.
     Example: "Negotiated with 15 external vendors during a supply chain disruption, reducing costs by 20%."
   - The same skill may appear in both sections, but the LANGUAGE must be completely different.
2. SENIORITY CALIBRATION:
   - If the target role is Entry/Mid-level (Buyer, Specialist, Coordinator, Analyst):
     Use operational verbs: Executed, Processed, Maintained, Resolved, Reconciled.
     Do NOT use: Orchestrated, Visionary, Spearheaded, Strategic Strategy.
   - If the target role is Senior (Director, VP, Head of):
     Use strategic verbs: Directed, Engineered, Optimized, Led.
3. HIDDEN REQUIREMENT DETECTION:
   - Scan the JD for soft skills (Customer Service, Communication, Analytical).
   - If the candidate data does not mention them, use Ghostwriter logic to generate a bullet.
4. Every bullet in Professional Experience must tie to at least one JD keyword.
   Quantify everything (dollars, percentages, personnel counts, timelines).
   Civilianize ALL military terms using the translation map.
5. If User Data is EMPTY or very thin, generate a "Top 10% Performer" resume from scratch
   based on Rank doctrine. Assume excellence: 100% accountability, zero loss, top ratings.
OUTPUT FORMAT (start immediately, no preamble, no "Here is the resume"):
{header_block}

## PROFESSIONAL SUMMARY
[3 sentences. Power statement: Rank Authority + top 3 JD keywords + Degree/Clearance.]

## CORE COMPETENCIES
* **[JD Keyword 1]:** [8-15 word context]
* **[JD Keyword 2]:** [context]
* **[JD Keyword 3]:** [context]
* **[JD Keyword 4]:** [context]
* **[JD Keyword 5]:** [context]
* **[Soft Skill from JD]:** [context]

## {proj_header}
* **[Project Name]:** [Action + quantified result matching a JD need]
* **[Project Name]:** [Action + quantified result matching a JD need]

## PROFESSIONAL EXPERIENCE
**[Civilianized Job Title]** (Former {rank}) | **U.S. Army**
*[Start Date] - [End Date]*
* [STAR bullet 1: Deep dive into top JD requirement. Different wording than Core Competencies.]
* [STAR bullet 2: Different JD requirement.]
* [STAR bullet 3: Leadership/mentorship with metrics.]
* [STAR bullet 4: Audit, compliance, or cost-savings metric.]
* [STAR bullet 5: Process improvement or customer service.]

## EDUCATION & CERTIFICATIONS
* [Degree] | [Institution] | [Year]
* [Clearance Level]
"""


def prompt_cover_letter(rank, years, industry, target_title, keywords, user_data, contact_info=None, gap_info=None, company_name=None):
    ctx = _context_block(rank, years, industry, target_title, keywords, user_data, contact_info, gap_info)
    company = company_name if company_name and company_name != "Unknown Company" else "[Company Name]"
    if contact_info and contact_info.get("name"):
        cl_name = contact_info["name"]
        cl_city = contact_info.get("city") or "[City, State]"
        cl_phone = contact_info.get("phone") or "[Phone]"
        cl_email = contact_info.get("email") or "[Email]"
        cl_header = f"{cl_name}\n{cl_city} | {cl_phone} | {cl_email}"
    else:
        cl_name = "[Full Name]"
        cl_header = "[Full Name]\n[City, State] | [Phone] | [Email]"
    today = datetime.now().strftime("%B %d, %Y")
    return f"""You are a Career Architect writing a cover letter for a 92Y veteran.
{ctx}
COMPANY: {company}
RULES:
- 3 paragraphs: Hook (who you are + why this role), Body (2-3 JD keywords matched to experience), Close (call to action).
- Address the letter to "{company}" hiring team.
- If there is an employment gap, frame it as a "Professional Development period" naturally.
- Civilianize all military terms. Match the industry tone.
- Under 350 words. Do NOT repeat the resume verbatim. Use different examples and angles.
OUTPUT (start immediately, no preamble):
{cl_header}

{today}

Dear {company} Hiring Team,

[Paragraph 1: Hook]

[Paragraph 2: Body]

[Paragraph 3: Close]

Respectfully,
{cl_name}
"""


def prompt_interview(rank, years, industry, target_title, keywords, user_data, contact_info=None, gap_info=None, company_name=None):
    ctx = _context_block(rank, years, industry, target_title, keywords, user_data, contact_info, gap_info)
    company = company_name if company_name and company_name != "Unknown Company" else "the company"
    return f"""You are an Interview Coach for a 92Y veteran applying to {company}.
{ctx}
Generate interview prep. Start immediately, no preamble.

## LIKELY INTERVIEW QUESTIONS

**Q1: [Question targeting top JD keyword]**
*Suggested Answer:* [STAR format, 3-4 sentences, using civilianized military experience.]

**Q2: [Question targeting second JD keyword]**
*Suggested Answer:* [STAR format.]

**Q3: [Behavioral question about leadership]**
*Suggested Answer:* [STAR format.]

**Q4: [Behavioral question about problem-solving or customer service]**
*Suggested Answer:* [STAR format.]

**Q5: [Industry-specific technical question]**
*Suggested Answer:* [STAR format.]

## QUESTIONS TO ASK THE INTERVIEWER
1. [Smart question demonstrating JD knowledge about {company}]
2. [Question about team structure or growth]
3. [Question about success metrics for the role]

## MILITARY TRANSLATION CHEAT SHEET
| If They Ask About... | Translate Your Experience As... |
|---|---|
| [Civilian concept 1] | [Military equivalent, civilianized] |
| [Civilian concept 2] | [Military equivalent, civilianized] |
| [Civilian concept 3] | [Military equivalent, civilianized] |
"""


def prompt_ats_analysis(resume_text, keywords):
    kw = "\n".join(f"  - {k}" for k in keywords)
    return f"""You are an ATS (Applicant Tracking System) analyst.
Analyze this resume against the target keywords and return a JSON report.
RESUME TEXT:
{resume_text[:5000]}
TARGET KEYWORDS:
{kw}
Return ONLY valid JSON. No preamble, no markdown fences:
{{
  "keyword_hits": {{
    "keyword_name": {{"found": true, "count": 2}},
    "other_keyword": {{"found": false, "count": 0}}
  }},
  "overall_density_score": <integer 0-100>,
  "missing_keywords": ["keyword1", "keyword2"],
  "suggestions": ["suggestion 1", "suggestion 2"]
}}
SCORING: 90-100 = excellent keyword coverage, 70-89 = good, 50-69 = needs work, below 50 = poor."""


# ============================================================
# 7. DOCX EXPORT
# ============================================================

def markdown_to_docx(md_text: str) -> io.BytesIO:
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10.5)
    style_normal.paragraph_format.space_after = Pt(2)
    style_normal.paragraph_format.space_before = Pt(0)
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# ") and not line.startswith("## "):
            text = line.lstrip("# ").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith("### "):
            text = re.sub(r"\*{1,2}", "", line.lstrip("# ").strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("## "):
            text = line.lstrip("# ").strip()
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:space"): "1", qn("w:color"): "1A1A2E",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif line.startswith("* ") or line.startswith("- "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, bullet_text)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
        elif "|" in line and ("@" in line or "phone" in line.lower() or "linkedin" in line.lower()):
            text = re.sub(r"\*{1,2}", "", line).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith("|") and line.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not all(c in "-| :" for c in row):
                    cells = [c.strip() for c in row.split("|")[1:-1]]
                    table_lines.append(cells)
                i += 1
            if table_lines:
                ncols = max(len(r) for r in table_lines)
                table = doc.add_table(rows=len(table_lines), cols=ncols)
                table.style = "Light Grid Accent 1"
                for ri, row_data in enumerate(table_lines):
                    for ci, cell_text in enumerate(row_data):
                        if ci < ncols:
                            cell = table.cell(ri, ci)
                            cell.text = re.sub(r"\*{1,2}", "", cell_text)
                            for par in cell.paragraphs:
                                for run in par.runs:
                                    run.font.size = Pt(9.5)
                                    run.font.name = "Calibri"
            continue
        elif line.startswith("**") and "**" in line[2:]:
            p = doc.add_paragraph()
            _add_runs(p, line)
            p.paragraph_format.space_after = Pt(1)
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            text = line.strip("*").strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(1)
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)
            p.paragraph_format.space_after = Pt(3)
        i += 1
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_runs(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


# ============================================================
# 8. ZIP EXPORT
# ============================================================

def create_zip_bundle(docx_files, company_slug, title_slug):
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for label, buf in docx_files.items():
            if buf and buf.getbuffer().nbytes > 0:
                zf.writestr(f"{label}_{company_slug}_{title_slug}.docx", buf.getvalue())
    zip_buf.seek(0)
    return zip_buf


# ============================================================
# 9. INPUT VALIDATION
# ============================================================

def validate_inputs(api_key, job_desc, target_title):
    if not api_key:
        return "API Key is required."
    if not job_desc or len(job_desc.split()) < 15:
        return "Job Description is too short. Paste the full JD (minimum ~15 words)."
    if not target_title or len(target_title.strip()) < 3:
        return "Target Job Title is required."
    return None


# ============================================================
# 10. DISPLAY HELPERS
# ============================================================

def display_match_score(score_data):
    score = score_data.get("score", 0)
    matched = score_data.get("matched", [])
    gaps = score_data.get("gaps", [])
    summary = score_data.get("summary", "")
    if score >= 80:
        color, label = "#28a745", "Strong Match"
    elif score >= 60:
        color, label = "#17a2b8", "Solid Match"
    elif score >= 40:
        color, label = "#ffc107", "Stretch Role"
    else:
        color, label = "#dc3545", "Significant Gaps"
    st.markdown(f"""
<div style="border: 2px solid {color}; border-radius: 10px; padding: 20px; margin: 10px 0;">
    <div style="display: flex; align-items: center; gap: 20px; margin-bottom: 12px;">
        <div style="font-size: 48px; font-weight: bold; color: {color};">{score}%</div>
        <div>
            <div style="font-size: 20px; font-weight: bold; color: {color};">{label}</div>
            <div style="font-size: 14px; color: #666;">{summary}</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
    mc1, mc2 = st.columns(2)
    with mc1:
        if matched:
            st.markdown("**Skills You Match:**")
            for m in matched:
                st.markdown(f"- {m}")
    with mc2:
        if gaps:
            st.markdown("**Gaps to Address:**")
            for g in gaps:
                st.markdown(f"- {g}")
    if score < 40:
        st.warning(
            "This role has significant gaps compared to your profile. "
            "You can still generate a resume, but consider upskilling in the gap areas "
            "or targeting a role that better fits your current experience."
        )
    elif score < 60:
        st.info(
            "This is a stretch role. The AI will use Ghostwriter logic to fill gaps with "
            "transferable skills from your rank, but you should prepare to address the "
            "gaps directly in interviews."
        )


def display_ats_analysis(ats_data):
    if not ats_data:
        return
    density = ats_data.get("overall_density_score", 0)
    hits = ats_data.get("keyword_hits", {})
    missing = ats_data.get("missing_keywords", [])
    suggestions = ats_data.get("suggestions", [])
    if density >= 90:
        color, label = "#28a745", "Excellent"
    elif density >= 70:
        color, label = "#17a2b8", "Good"
    elif density >= 50:
        color, label = "#ffc107", "Needs Work"
    else:
        color, label = "#dc3545", "Poor"
    st.markdown(f"""
<div style="border: 2px solid {color}; border-radius: 8px; padding: 15px; margin: 10px 0;">
    <div style="display: flex; align-items: center; gap: 15px;">
        <div style="font-size: 36px; font-weight: bold; color: {color};">{density}%</div>
        <div>
            <div style="font-size: 16px; font-weight: bold; color: {color};">ATS Keyword Density: {label}</div>
            <div style="font-size: 13px; color: #666;">How well your resume matches the job description keywords</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)
    if hits:
        found_count = sum(1 for v in hits.values() if v.get("found"))
        total = len(hits)
        st.markdown(f"**Keyword Coverage:** {found_count}/{total} keywords found in resume")
    if missing:
        st.markdown("**Missing Keywords:** " + ", ".join(missing))
        st.caption("Consider adding these to your resume if you have relevant experience.")
    if suggestions:
        with st.expander("Optimization Suggestions"):
            for s in suggestions:
                st.markdown(f"- {s}")


def display_next_steps(company_name, target_title):
    company = company_name if company_name and company_name != "Unknown Company" else "the company"
    st.markdown("### What To Do Next")
    st.markdown(f"""
**1. Submit your application** through {company}'s careers page or the job board where you found the listing.

**2. Find the hiring manager or recruiter** for this role on LinkedIn. Search for "{target_title}" + "{company}" or look at the company's People page filtered by "Recruiting" or "Talent Acquisition."

**3. Send a connection request** with a short note:
> "Hi [Name], I recently applied for the {target_title} position and wanted to connect. My background in supply chain operations and asset management aligns well with the role. I'd welcome the chance to discuss how I can contribute to the team."

**4. Follow up once** after 5-7 business days if you don't hear back. Keep it brief, reference your application date, and restate your interest.

**5. Prep for the interview** using the Interview Prep tab above. Practice your answers out loud at least twice before the real thing.
""")


# ============================================================
# 11. UI LAYOUT
# ============================================================

st.title("92Y Career Auto-Pilot")
st.markdown("##### Military-to-Civilian Resume Engine with ATS Keyword Matching")

# Sidebar
with st.sidebar:
    st.header("Authorization")
    default_key = ""
    try:
        default_key = st.secrets.get("GOOGLE_API_KEY", "")
    except Exception:
        pass
    if default_key:
        api_key = default_key
        st.success("API key loaded from server.")
    else:
        api_key = st.text_input("Google Gemini API Key", type="password")
        st.caption("Free key at [aistudio.google.com](https://aistudio.google.com/apikey)")
    st.divider()
    st.markdown(
        "**Generates:**\n"
        "1. Job Match Score\n"
        "2. ATS-optimized Resume (.docx)\n"
        "3. Tailored Cover Letter (.docx)\n"
        "4. Interview Prep Guide (.docx)\n"
        "5. ATS Keyword Analysis\n"
        "6. Next Steps Action Plan"
    )
    st.divider()
    st.caption(
        "v5.0 | Mirror Protocol | Ghostwriter | "
        "No-Repeat Rule | 4-Industry Logic | "
        "Seniority Calibration | Match Score | "
        "Chameleon Headers | ATS Analysis | "
        "Contact Info | Gap Handler | .docx Export"
    )

# ============================================================
# 12. CONTACT INFO (collapsible)
# ============================================================

with st.expander("Your Contact Information (recommended for a polished resume)"):
    st.caption("Fill in what you have. Anything left blank will show as a placeholder you can edit later in the .docx file.")
    ci1, ci2 = st.columns(2)
    with ci1:
        contact_name = st.text_input("Full Name", placeholder="e.g., James Rodriguez")
        contact_email = st.text_input("Email", placeholder="e.g., james.rodriguez@email.com")
        contact_linkedin = st.text_input("LinkedIn URL (optional)", placeholder="e.g., linkedin.com/in/jrodriguez")
    with ci2:
        contact_city = st.text_input("City, State", placeholder="e.g., Augusta, GA")
        contact_phone = st.text_input("Phone", placeholder="e.g., (706) 555-1234")

contact_info = {
    "name": contact_name.strip() if contact_name else "",
    "email": contact_email.strip() if contact_email else "",
    "city": contact_city.strip() if contact_city else "",
    "phone": contact_phone.strip() if contact_phone else "",
    "linkedin": contact_linkedin.strip() if contact_linkedin else "",
}
if not contact_info["name"]:
    contact_info = None

# ============================================================
# 13. EMPLOYMENT GAP (collapsible)
# ============================================================

with st.expander("Employment Gap? (optional)"):
    st.caption("If you have a gap between military service and now, fill this in. The AI will frame it positively.")
    has_gap = st.checkbox("I have an employment gap")
    gap_info = None
    if has_gap:
        g1, g2 = st.columns(2)
        with g1:
            gap_start = st.text_input("Gap Start", placeholder="e.g., March 2022")
        with g2:
            gap_end = st.text_input("Gap End", value="Present", placeholder="e.g., Present")
        gap_activities = st.multiselect(
            "What did you do during the gap?",
            options=[
                "Completed Bachelor's Degree",
                "Completed Master's Degree",
                "Pursued Certifications (PMP, CSCP, etc.)",
                "Freelance/Contract Work",
                "Volunteered",
                "Family Caregiving",
                "Skills Training / Bootcamp",
                "Started a Business",
                "Relocated",
            ],
            default=["Pursued Certifications (PMP, CSCP, etc.)"],
        )
        gap_other = st.text_input("Other activities (optional):", placeholder="e.g., Completed SFL-TAP program")
        all_activities = list(gap_activities)
        if gap_other.strip():
            all_activities.append(gap_other.strip())
        gap_info = {
            "has_gap": True,
            "start": gap_start,
            "end": gap_end,
            "activities": all_activities,
        }

# ============================================================
# 14. PROFILE & TARGET INPUTS
# ============================================================

col1, col2 = st.columns(2)

with col1:
    st.subheader("Your Profile")
    c1, c2 = st.columns(2)
    with c1:
        rank = st.selectbox("Rank", ["E-4 (SPC)", "E-5 (SGT)", "E-6 (SSG)", "E-7 (SFC)", "E-8 (MSG)"])
    with c2:
        years = st.number_input("Years of Service", 1, 30, 9)
    input_tab1, input_tab2, input_tab3 = st.tabs(
        ["Upload Resume", "Paste Bullets", "Generate from Scratch"]
    )
    user_data = ""
    with input_tab1:
        uploaded_file = st.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])
        if uploaded_file:
            try:
                user_data = read_file(uploaded_file)
                st.success(f"Loaded {len(user_data.split())} words from {uploaded_file.name}")
            except (ValueError, Exception) as e:
                st.error(str(e))
    with input_tab2:
        manual_text = st.text_area(
            "Paste NCOER bullets, award citations, or brain dump:",
            height=200,
            placeholder="- Maintained 100% accountability of $15M property book\n- Scored 98% on CSDP inspection\n- Processed 15 FLIPLs recovering $200k",
        )
        if manual_text:
            user_data = manual_text
    with input_tab3:
        st.markdown(
            "**No resume or bullets?** Select your Rank and Years above, fill in the "
            "Target Position on the right, and the AI will build everything from scratch "
            "based on standard 92Y duties for your rank. It assumes you were a top performer."
        )
        scratch_confirm = st.checkbox("I want to generate from scratch (no upload needed)")
        if scratch_confirm and not user_data:
            user_data = ""

with col2:
    st.subheader("Target Position")
    target_ind = st.selectbox(
        "Target Industry",
        ["Corporate (General)", "Defense Contractor", "Federal (USAJOBS)", "Tech / SaaS"],
    )
    target_title = st.text_input("Target Job Title", placeholder="e.g., Procurement Buyer II")
    resume_pages = st.selectbox(
        "Resume Length",
        ["1 page", "2 pages (recommended)", "3+ pages (Federal only)"],
        index=1,
    )
    job_desc = st.text_area(
        "Paste Full Job Description",
        height=200,
        placeholder="Paste the complete JD here. The AI will scan for every requirement.",
    )


# ============================================================
# 15. STEP 1: KEYWORD EXTRACTION + MATCH SCORE
# ============================================================

st.divider()

if st.button("Step 1: Extract JD Keywords", type="secondary"):
    err = validate_inputs(api_key, job_desc, target_title)
    if err:
        st.warning(err)
    else:
        try:
            with st.spinner("Scanning job description..."):
                if not st.session_state.model:
                    st.session_state.model = init_model(api_key)
                raw = call_model(st.session_state.model, prompt_keywords(job_desc))
                kws = json.loads(raw)
                if not isinstance(kws, list) or len(kws) < 3:
                    raise ValueError("Too few keywords returned.")
                st.session_state.keywords = kws
                try:
                    company = call_model(st.session_state.model, prompt_company_extract(job_desc))
                    company = company.strip().strip('"').strip("'")
                    st.session_state.company_name = company if company else "Unknown Company"
                except Exception:
                    st.session_state.company_name = "Unknown Company"
                score_raw = call_model(
                    st.session_state.model,
                    prompt_match_score(rank, years, kws, user_data, target_title),
                )
                score_data = json.loads(score_raw)
                st.session_state.match_score = score_data
                st.session_state.keywords_confirmed = False
                st.session_state.generation_complete = False
                st.session_state.resume_md = None
                st.session_state.cover_letter_md = None
                st.session_state.interview_md = None
                st.session_state.ats_analysis = None
                st.session_state.generated_at = None
                st.rerun()
        except json.JSONDecodeError:
            st.error("Failed to parse AI response. Try again.")
        except Exception as e:
            st.error(f"Analysis failed: {e}")


# ============================================================
# 16. SHOW MATCH SCORE + EDITABLE KEYWORDS
# ============================================================

if st.session_state.keywords and not st.session_state.keywords_confirmed:
    if st.session_state.match_score:
        display_match_score(st.session_state.match_score)
        st.divider()
    if st.session_state.company_name and st.session_state.company_name != "Unknown Company":
        st.markdown(f"**Company Detected:** {st.session_state.company_name}")
    st.markdown("**Extracted JD Keywords** (edit, remove, or add):")
    edited = []
    cols = st.columns(3)
    for idx, kw in enumerate(st.session_state.keywords):
        with cols[idx % 3]:
            val = st.text_input(f"Keyword {idx+1}", value=kw, key=f"kw_{idx}")
            if val.strip():
                edited.append(val.strip())
    add_kw = st.text_input("Add a keyword (optional):", key="add_kw", placeholder="e.g., Lean Six Sigma")
    kc1, kc2 = st.columns(2)
    with kc1:
        if st.button("Confirm Keywords & Generate", type="primary"):
            if add_kw.strip():
                edited.append(add_kw.strip())
            st.session_state.keywords = edited
            st.session_state.keywords_confirmed = True
            st.rerun()
    with kc2:
        if st.button("Re-extract"):
            st.session_state.keywords = None
            st.session_state.match_score = None
            st.session_state.company_name = None
            st.rerun()


# ============================================================
# 17. STEP 2: GENERATE ALL SECTIONS (error recovery)
# ============================================================

if st.session_state.keywords_confirmed and not st.session_state.generation_complete:
    model = st.session_state.model
    kws = st.session_state.keywords
    company = st.session_state.company_name
    if not user_data:
        user_data = f"[NO DATA PROVIDED. Generate from scratch for {rank} with {years} years of 92Y service. Assume top 10% performer.]"
    progress = st.progress(0, text="Starting generation...")
    errors = []
    try:
        progress.progress(10, text="Generating resume...")
        st.session_state.resume_md = call_model(
            model, prompt_resume(rank, years, target_ind, target_title, kws, user_data, contact_info, gap_info, resume_pages)
        )
    except Exception as e:
        errors.append(f"Resume: {e}")
        st.session_state.resume_md = None
    try:
        progress.progress(35, text="Generating cover letter...")
        st.session_state.cover_letter_md = call_model(
            model, prompt_cover_letter(rank, years, target_ind, target_title, kws, user_data, contact_info, gap_info, company)
        )
    except Exception as e:
        errors.append(f"Cover Letter: {e}")
        st.session_state.cover_letter_md = None
    try:
        progress.progress(60, text="Generating interview prep...")
        st.session_state.interview_md = call_model(
            model, prompt_interview(rank, years, target_ind, target_title, kws, user_data, contact_info, gap_info, company)
        )
    except Exception as e:
        errors.append(f"Interview Prep: {e}")
        st.session_state.interview_md = None
    if st.session_state.resume_md:
        try:
            progress.progress(80, text="Running ATS keyword analysis...")
            ats_raw = call_model(model, prompt_ats_analysis(st.session_state.resume_md, kws))
            st.session_state.ats_analysis = json.loads(ats_raw)
        except Exception:
            st.session_state.ats_analysis = None
    progress.progress(100, text="Complete!")
    if st.session_state.resume_md or st.session_state.cover_letter_md or st.session_state.interview_md:
        st.session_state.generation_complete = True
        st.session_state.generated_at = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    else:
        st.error("All sections failed to generate. Check your API key and try again.")
    if errors:
        for err in errors:
            st.warning(f"Partial failure - {err}")
    time.sleep(0.5)
    if st.session_state.generation_complete:
        st.rerun()


# ============================================================
# 18. DISPLAY RESULTS
# ============================================================

if st.session_state.generation_complete:
    st.divider()
    st.subheader("Your Career Package")
    meta_parts = []
    if st.session_state.company_name and st.session_state.company_name != "Unknown Company":
        meta_parts.append(f"**Company:** {st.session_state.company_name}")
    if st.session_state.match_score:
        meta_parts.append(f"**Match Score:** {st.session_state.match_score.get('score', 'N/A')}%")
    if st.session_state.generated_at:
        meta_parts.append(f"**Generated:** {st.session_state.generated_at}")
    if meta_parts:
        st.markdown(" | ".join(meta_parts))
    company_slug = "Resume"
    if st.session_state.company_name and st.session_state.company_name != "Unknown Company":
        company_slug = re.sub(r"[^a-zA-Z0-9]", "_", st.session_state.company_name)
    title_slug = re.sub(r"[^a-zA-Z0-9]", "_", target_title) if target_title else "Role"

    # ATS Analysis
    if st.session_state.ats_analysis:
        with st.expander("ATS Keyword Analysis", expanded=True):
            display_ats_analysis(st.session_state.ats_analysis)

    # Build tabs dynamically based on what succeeded
    available_tabs = []
    tab_labels = []
    if st.session_state.resume_md:
        available_tabs.append("resume")
        tab_labels.append("Resume")
    if st.session_state.cover_letter_md:
        available_tabs.append("cover_letter")
        tab_labels.append("Cover Letter")
    if st.session_state.interview_md:
        available_tabs.append("interview")
        tab_labels.append("Interview Prep")
    tabs = st.tabs(tab_labels)
    docx_files = {}
    for idx, tab_key in enumerate(available_tabs):
        with tabs[idx]:
            if tab_key == "resume":
                st.markdown(st.session_state.resume_md)
                resume_docx = markdown_to_docx(st.session_state.resume_md)
                docx_files["Resume"] = resume_docx
                st.download_button(
                    "Download Resume (.docx)",
                    data=resume_docx,
                    file_name=f"Resume_{company_slug}_{title_slug}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            elif tab_key == "cover_letter":
                st.markdown(st.session_state.cover_letter_md)
                cl_docx = markdown_to_docx(st.session_state.cover_letter_md)
                docx_files["Cover_Letter"] = cl_docx
                st.download_button(
                    "Download Cover Letter (.docx)",
                    data=cl_docx,
                    file_name=f"Cover_Letter_{company_slug}_{title_slug}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            elif tab_key == "interview":
                st.markdown(st.session_state.interview_md)
                int_docx = markdown_to_docx(st.session_state.interview_md)
                docx_files["Interview_Prep"] = int_docx
                st.download_button(
                    "Download Interview Prep (.docx)",
                    data=int_docx,
                    file_name=f"Interview_Prep_{company_slug}_{title_slug}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

    # Download All
    if len(docx_files) > 1:
        st.divider()
        zip_bundle = create_zip_bundle(docx_files, company_slug, title_slug)
        st.download_button(
            "Download All (.zip)",
            data=zip_bundle,
            file_name=f"Career_Package_{company_slug}_{title_slug}.zip",
            mime="application/zip",
        )

    # Next Steps
    display_next_steps(st.session_state.company_name, target_title)

    # Start Over
    st.markdown("---")
    if st.button("Start Over"):
        for key in STATE_DEFAULTS:
            st.session_state[key] = STATE_DEFAULTS[key]
        st.rerun()
